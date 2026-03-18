from __future__ import annotations

import io
import json
import os
from datetime import datetime
from pathlib import Path

from flask import Flask, flash, redirect, render_template, request, send_file, url_for
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt, RGBColor

BASE_DIR = Path(__file__).resolve().parent
DATA_DIR = BASE_DIR / "data"
EXPORTS_DIR = BASE_DIR / "exports"
ENTRIES_FILE = DATA_DIR / "entries.json"

DATA_DIR.mkdir(exist_ok=True)
EXPORTS_DIR.mkdir(exist_ok=True)

app = Flask(__name__)
app.config["SECRET_KEY"] = os.getenv("SECRET_KEY", "worddist-secret-change-me")


def load_entries() -> list[dict]:
    if not ENTRIES_FILE.exists():
        return []

    with ENTRIES_FILE.open("r", encoding="utf-8") as f:
        data = json.load(f)

    if not isinstance(data, list):
        return []

    return data


def save_entries(entries: list[dict]) -> None:
    with ENTRIES_FILE.open("w", encoding="utf-8") as f:
        json.dump(entries, f, ensure_ascii=False, indent=2)


def build_word(entries: list[dict]) -> io.BytesIO:
    doc = Document()

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title.add_run("Acuerdo del Equipo de Padel")
    title_run.bold = True
    title_run.font.size = Pt(20)
    title_run.font.color.rgb = RGBColor(24, 76, 142)

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle_run = subtitle.add_run(
        f"Generado por WordDist | {datetime.now().strftime('%d/%m/%Y %H:%M')}"
    )
    subtitle_run.italic = True
    subtitle_run.font.size = Pt(10)

    doc.add_paragraph("")

    rights = [e for e in entries if e.get("entry_type") == "derecho"]
    duties = [e for e in entries if e.get("entry_type") == "obligacion"]

    for section_title, section_data in (
        ("Derechos de los jugadores", rights),
        ("Obligaciones de los jugadores", duties),
    ):
        heading = doc.add_paragraph()
        heading_run = heading.add_run(section_title)
        heading_run.bold = True
        heading_run.font.size = Pt(14)
        heading_run.font.color.rgb = RGBColor(30, 30, 30)

        if not section_data:
            doc.add_paragraph("- Sin entradas todavia")
            doc.add_paragraph("")
            continue

        for idx, item in enumerate(section_data, start=1):
            author = item.get("author") or "Anonimo"
            text = item.get("text", "").strip()
            created_at = item.get("created_at", "")

            paragraph = doc.add_paragraph(style="List Number")
            paragraph.add_run(f"{text} ")

            meta = paragraph.add_run(f"(propuesto por {author} - {created_at})")
            meta.italic = True
            meta.font.size = Pt(9)
            meta.font.color.rgb = RGBColor(90, 90, 90)

        doc.add_paragraph("")

    footer = doc.add_paragraph()
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer_run = footer.add_run("WordDist - Documento colaborativo del equipo")
    footer_run.font.size = Pt(9)
    footer_run.font.color.rgb = RGBColor(120, 120, 120)

    out = io.BytesIO()
    doc.save(out)
    out.seek(0)
    return out


@app.route("/", methods=["GET"])
def dashboard():
    entries = sorted(load_entries(), key=lambda x: x.get("created_at", ""), reverse=True)
    return render_template("index.html", entries=entries)


@app.route("/opinar", methods=["GET", "POST"])
def opinar():
    if request.method == "POST":
        entry_type = request.form.get("entry_type", "").strip().lower()
        author = request.form.get("author", "").strip()
        text = request.form.get("text", "").strip()

        if entry_type not in {"derecho", "obligacion"}:
            flash("Selecciona si tu propuesta es derecho u obligacion.", "error")
            return redirect(url_for("opinar"))

        if not text:
            flash("La propuesta no puede quedar vacia.", "error")
            return redirect(url_for("opinar"))

        entries = load_entries()
        entries.append(
            {
                "entry_type": entry_type,
                "author": author,
                "text": text,
                "created_at": datetime.now().strftime("%d/%m/%Y %H:%M"),
            }
        )
        save_entries(entries)

        flash("Entrada guardada correctamente. Gracias por participar.", "ok")
        return redirect(url_for("opinar"))

    return render_template("opinar.html")


@app.route("/exportar-word", methods=["GET"])
def exportar_word():
    entries = load_entries()
    file_obj = build_word(entries)

    filename = f"worddist_equipo_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx"
    return send_file(
        file_obj,
        as_attachment=True,
        download_name=filename,
        mimetype="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    )


if __name__ == "__main__":
    port = int(os.getenv("PORT", "5000"))
    debug = os.getenv("FLASK_DEBUG", "false").lower() == "true"
    app.run(host="0.0.0.0", port=port, debug=debug)
